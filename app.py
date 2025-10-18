import streamlit as st
import os
from typing import List
import re
import base64
import fitz  # PyMuPDF
import docx
import requests

def inject_fixed_header():
    st.markdown("""
    <style>
    div[data-testid="stToolbar"] {visibility: hidden;}
    div[data-testid="stDecoration"] {visibility: hidden;}
    div[data-testid="stStatusWidget"] {visibility: hidden;}
    #MainMenu {visibility: hidden;}
    header {visibility: hidden;}

    /* ---- Fixed header ---- */
    .netsec-fixed-header {
        position: fixed !important;
        top: 0;
        left: 0;
        width: 100%;
        background: #0e1117;
        text-align: center;
        z-index: 999999;
        border-bottom: 1px solid rgba(255,255,255,0.15);
        padding: 18px 0 12px 0;
    }

    .netsec-fixed-header h1 {
        color: white !important;
        font-weight: 800;
        font-size: 32px;
        margin-bottom: 4px;
        letter-spacing: 0.5px;
        font-family: "Segoe UI", sans-serif;
    }

    .netsec-fixed-header p {
        color: #ccc !important;
        font-size: 17px;
        margin: 0;
        opacity: 0.85;
    }

    /* Prevent overlap with header */
    [data-testid="stAppViewContainer"] .block-container {
        padding-top: 120px !important;
    }
    </style>

    <div class="netsec-fixed-header">
        <h1>🔐 NetSec Tutor</h1>
        <p>An easy way to study and test your network security skills…</p>
    </div>
    """, unsafe_allow_html=True)


# --- Summarization helpers (module level) ---
def handle_summarization_request(prompt: str):
    """
    Improved: Robust file matching, always returns summary + citation dict for chat history.
    """
    # Patterns for textbook and slides
    page_pat = re.compile(r"summari[sz]e\s+page\s+(\d+)\s+from\s+([\w\d_\-\.]+)", re.I)
    slide_pat = re.compile(r"summari[sz]e\s+slide(?:\s*no)?\s*(\d+)\s+from\s+lecture slides?\s*([\w\d_\-\.]+)", re.I)
    m_page = page_pat.search(prompt)
    m_slide = slide_pat.search(prompt)
    kb_files = _get_kb_files()
    if m_page:
        page_num = int(m_page.group(1))
        doc_name = m_page.group(2).lower()
        # Try to match by index (e.g. 'textbook1' means first file)
        files = kb_files.get("TEXTBOOKs", [])
        file_path = None
        if doc_name.startswith("textbook") and doc_name[8:].isdigit():
            idx = int(doc_name[8:]) - 1
            if 0 <= idx < len(files):
                file_path = files[idx]
        # Fallback: match by normalized name
        if not file_path:
            for f in files:
                if doc_name in os.path.splitext(os.path.basename(f))[0].lower():
                    file_path = f
                    break
        if not file_path:
            return {"error": f"No textbook found matching '{doc_name}'."}
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        if ext == ".pdf":
            try:
                doc = fitz.open(file_path)
                if page_num < 1 or page_num > doc.page_count:
                    return {"error": f"Page {page_num} out of range (1-{doc.page_count}) for {filename}."}
                text = doc.load_page(page_num-1).get_text()
            except Exception as e:
                return {"error": f"Error reading PDF: {e}"}
        elif ext == ".docx":
            try:
                doc = docx.Document(file_path)
                text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                return {"error": f"Error reading DOCX: {e}"}
        elif ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            except Exception as e:
                return {"error": f"Error reading TXT: {e}"}
        else:
            return {"error": f"Unsupported file type: {ext}"}
        summary = ask_llm_summary(text, page_num, filename)
        citation = {
            "filename": filename,
            "page_number": page_num,
            "content_preview": text[:300] + ("..." if len(text) > 300 else ""),
            "full_text": text
        }
        return {"summary": summary, "citation": citation}
    elif m_slide:
        slide_num = int(m_slide.group(1))
        slide_name = m_slide.group(2).lower()
        files = kb_files.get("LECTURE SLIDEs", [])
        file_path = None
        if slide_name.isdigit():
            idx = int(slide_name) - 1
            if 0 <= idx < len(files):
                file_path = files[idx]
        if not file_path:
            for f in files:
                if slide_name in os.path.splitext(os.path.basename(f))[0].lower():
                    file_path = f
                    break
        if not file_path:
            return {"error": f"No lecture slide found matching '{slide_name}'."}
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        if ext == ".pdf":
            try:
                doc = fitz.open(file_path)
                if slide_num < 1 or slide_num > doc.page_count:
                    return {"error": f"Slide {slide_num} out of range (1-{doc.page_count}) for {filename}."}
                text = doc.load_page(slide_num-1).get_text()
            except Exception as e:
                return {"error": f"Error reading PDF: {e}"}
        elif ext == ".docx":
            try:
                doc = docx.Document(file_path)
                text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                return {"error": f"Error reading DOCX: {e}"}
        elif ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            except Exception as e:
                return {"error": f"Error reading TXT: {e}"}
        else:
            return {"error": f"Unsupported file type: {ext}"}
        summary = ask_llm_summary(text, slide_num, filename, is_slide=True)
        citation = {
            "filename": filename,
            "page_number": slide_num,
            "content_preview": text[:300] + ("..." if len(text) > 300 else ""),
            "full_text": text
        }
        return {"summary": summary, "citation": citation}
    else:
        return None

def ask_llm_summary_stream(text, num, filename, is_slide=False):
    """
    Streams summary tokens from Ollama API (yields tokens as they arrive).
    """
    prompt = f"Summarize the following {'slide' if is_slide else 'page'} (number {num}) from {filename}:\n\n{text}\n\nSummary:"
    url = f"{Config.OLLAMA_BASE_URL}/api/generate"
    try:
        with requests.post(
            url,
            json={
                "model": Config.OLLAMA_MODEL,
                "prompt": prompt,
                "temperature": Config.OLLAMA_TEMPERATURE,
                "stream": True
            },
            timeout=120,
            stream=True
        ) as resp:
            if resp.status_code != 200:
                yield {"type": "error", "text": f"LLM error: {resp.text}"}
                return
            buffer = ""
            for line in resp.iter_lines():
                if not line:
                    continue
                try:
                    data = line.decode("utf-8")
                    # Ollama streams JSON objects per line
                    import json as _json
                    obj = _json.loads(data)
                    token = obj.get("response", "")
                    if token:
                        buffer += token
                        yield {"type": "token", "text": token, "buffer": buffer}
                    if obj.get("done"):
                        break
                except Exception:
                    continue
            yield {"type": "final", "text": buffer}
    except Exception as e:
        yield {"type": "error", "text": f"Error calling LLM: {e}"}

def ask_llm_summary(text, num, filename, is_slide=False):
    # Non-streaming fallback for summary
    prompt = f"Summarize the following {'slide' if is_slide else 'page'} (number {num}) from {filename}:\n\n{text}\n\nSummary:"
    try:
        response = requests.post(
            f"{Config.OLLAMA_BASE_URL}/api/generate",
            json={
                "model": Config.OLLAMA_MODEL,
                "prompt": prompt,
                "temperature": Config.OLLAMA_TEMPERATURE,
                "stream": False
            },
            timeout=60
        )
        if response.status_code == 200:
            data = response.json()
            return data.get("response", "[No summary returned]")
        else:
            return f"LLM error: {response.text}"
    except Exception as e:
        return f"Error calling LLM: {e}"
from utils.document_processor import DocumentProcessor
from utils.vector_store import VectorStore
from utils.rag_chain import RAGChain, CombinedRAGChain
from utils.session_store import (
    list_sessions,
    load_session,
    save_session,
    delete_session,
    create_new_session,
    most_recent_session,
)
from config import Config

# ------------------ Main Page UI Config ------------------
st.set_page_config(
    page_title=Config.PAGE_TITLE,
    page_icon=Config.PAGE_ICON,
    layout="wide"
)

# ------------------ Custom Styling using CSS with "unsafe_allow_html=True" ------------------
st.markdown("""
<style>
html, body, [class*="css"] {
    font-family: 'Segoe UI', sans-serif;
    overflow: hidden; /* prevent whole-page scroll */
}

/* ===== Layout adjustments ===== */
[data-testid="stAppViewContainer"] .block-container {
    padding-top: 120px !important; /* Space for header */
    padding-bottom: 100px !important;
    overflow: hidden;
}

/* ===== Scrollable chat area only ===== */
div[data-testid="stVerticalBlock"]:has(.stChatMessage) {
    max-height: calc(100vh - 220px);
    overflow-y: auto;
    padding-right: 12px;
    scroll-behavior: smooth;
}

/* ===== Fixed chat input ===== */
[data-testid="stChatInput"] {
    position: fixed !important;
    bottom: 0 !important;
    left: var(--chat-left-offset, 300px) !important;
    right: 0 !important;
    z-index: 1000 !important;
    background: var(--background-color, #0e1117);
    padding: 8px 16px 16px 16px;
    box-shadow: 0 -2px 8px rgba(0,0,0,0.1);
}

/* ===== Mobile sidebar collapse ===== */
@media (max-width: 900px) {
    [data-testid="stChatInput"] {
        left: 0 !important;
    }
}

/* ===== Smooth auto-scroll ===== */
</style>

<script>
const chatContainerObserver = new MutationObserver(() => {
    const chatBlocks = document.querySelectorAll('div[data-testid="stVerticalBlock"]');
    chatBlocks.forEach(block => {
        if (block.querySelector('.stChatMessage')) {
            block.scrollTop = block.scrollHeight;
        }
    });
});
chatContainerObserver.observe(document.body, { childList: true, subtree: true });
</script>
<style>
/* Hide Streamlit's built-in 'Limit 200MB per file' text */
[data-testid="stFileUploaderDropzone"] small {
    visibility: hidden;
}

/* Inject our own limit text */
[data-testid="stFileUploaderDropzone"]::after {
    content: "Limit 25 MB per file • PDF, DOCX, TXT";
    display: block;
    font-size: 0.8rem;
    color: rgba(255,255,255,0.7);
    text-align: center;
    margin-top: 4px;
}
</style>
""", unsafe_allow_html=True)


# ------------------ Initialize Session State ------------------
# Vectorstores for separate knowledge bases
if "vectorstore_textbooks" not in st.session_state:
    st.session_state.vectorstore_textbooks = None
if "vectorstore_slides" not in st.session_state:
    st.session_state.vectorstore_slides = None
if "rag_chain_textbooks" not in st.session_state:
    st.session_state.rag_chain_textbooks = None
if "rag_chain_slides" not in st.session_state:
    st.session_state.rag_chain_slides = None
if "rag_chain_both" not in st.session_state:
    st.session_state.rag_chain_both = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "documents_processed" not in st.session_state:
    st.session_state.documents_processed = False
# Chat session persistence
if "session_id" not in st.session_state:
    st.session_state.session_id = None
if "session_title" not in st.session_state:
    st.session_state.session_title = None
if "session_loaded" not in st.session_state:
    st.session_state.session_loaded = False
if "ignore_next_session_select" not in st.session_state:
    st.session_state.ignore_next_session_select = False

# ------------------ Helpers ------------------
def _get_kb_files() -> dict:
    files = {"TEXTBOOKs": [], "LECTURE SLIDEs": []}
    dirs = {
        "TEXTBOOKs": getattr(Config, "DATA_DIR_TEXTBOOKS", Config.DATA_DIR),
        "LECTURE SLIDEs": getattr(Config, "DATA_DIR_SLIDES", Config.DATA_DIR),
    }
    for label, kb_dir in dirs.items():
        try:
            if os.path.isdir(kb_dir):
                for name in sorted(os.listdir(kb_dir)):
                    path = os.path.join(kb_dir, name)
                    if os.path.isfile(path) and os.path.splitext(name)[1].lower() in {'.pdf', '.docx', '.txt'}:
                        files[label].append(path)
        except Exception:
            pass
    return files

def _render_kb_listing():
    files_by_cat = _get_kb_files()
    st.subheader("📚 Knowledge Base Files")
    empty = not (files_by_cat.get("TEXTBOOKs") or files_by_cat.get("LECTURE SLIDEs"))
    if empty:
        st.info("No documents in the knowledge base yet. Upload to get started.")
        return
    for cat_label, files in files_by_cat.items():
        st.markdown(f"**{cat_label}**")
        if not files:
            st.caption("(no files)")
            continue
        for path in files:
            name = os.path.basename(path)
            size = os.path.getsize(path)
            label = f"📄 {name} ({size/1024:.1f} KB)"
            with st.expander(label):
                try:
                    with open(path, 'rb') as fh:
                        st.download_button("⬇️ Download", data=fh.read(), file_name=name, use_container_width=True)
                except Exception:
                    st.warning("Download unavailable for this file.")

                # View button copies the file path to clipboard
                file_path = os.path.abspath(path)
                view_btn = st.button("👁️ Copy Path", key=f"view_{cat_label}_{name}", use_container_width=True)
                if view_btn:
                    st.code(file_path, language="text")
                    st.success("File path copied! Use your system's copy shortcut if needed.")


def initialize_components():
    # Create both vector stores
    if st.session_state.vectorstore_textbooks is None:
        st.session_state.vectorstore_textbooks = VectorStore(
            embedding_model=Config.EMBEDDING_MODEL,
            persist_directory=getattr(Config, "VECTOR_STORE_DIR_TEXTBOOKS", Config.VECTOR_STORE_DIR)
        )
        try:
            loaded = st.session_state.vectorstore_textbooks.load_vectorstore()
            if loaded:
                info = st.session_state.vectorstore_textbooks.get_collection_info()
                if info.get("count", 0) > 0:
                    st.session_state.documents_processed = True
        except Exception:
            pass
    if st.session_state.vectorstore_slides is None:
        st.session_state.vectorstore_slides = VectorStore(
            embedding_model=Config.EMBEDDING_MODEL,
            persist_directory=getattr(Config, "VECTOR_STORE_DIR_SLIDES", Config.VECTOR_STORE_DIR)
        )
        try:
            loaded = st.session_state.vectorstore_slides.load_vectorstore()
            if loaded:
                info = st.session_state.vectorstore_slides.get_collection_info()
                if info.get("count", 0) > 0:
                    st.session_state.documents_processed = True
        except Exception:
            pass

    # Initialize RAG chains for each store
    if st.session_state.rag_chain_textbooks is None and st.session_state.vectorstore_textbooks is not None:
        st.session_state.rag_chain_textbooks = RAGChain(st.session_state.vectorstore_textbooks)
    if st.session_state.rag_chain_slides is None and st.session_state.vectorstore_slides is not None:
        st.session_state.rag_chain_slides = RAGChain(st.session_state.vectorstore_slides)
    # Initialize combined RAG if possible
    if (
        st.session_state.rag_chain_both is None and 
        (st.session_state.vectorstore_textbooks is not None or st.session_state.vectorstore_slides is not None)
    ):
        stores = [s for s in [st.session_state.vectorstore_textbooks, st.session_state.vectorstore_slides] if s is not None]
        st.session_state.rag_chain_both = CombinedRAGChain(stores)

    # Default selections for upload and query categories
    st.session_state.setdefault("query_kb", "TEXTBOOKs")



# ------------------ Main UI ------------------
def main():
    inject_fixed_header()

    # Checking Ollama availability
    try:
        import requests
        response = requests.get(f"{Config.OLLAMA_BASE_URL}/api/tags")
        if response.status_code != 200:
            st.error("Could Not find any model from Ollama is running, make sure a model from Ollama is running.")
            st.info("Run ollama by 'ollama run mistral' ")
            return
    except Exception:
        st.error("Could Not find any model from Ollama is running, make sure a model from Ollama is running.")
        st.info("Run ollama by 'ollama run mistral' ")
        return

    # Initialize components
    initialize_components()

    # Load most recent session if any
    if not st.session_state.session_loaded:
        recent = most_recent_session()
        if recent:
            st.session_state.session_id = recent.get("session_id")
            st.session_state.session_title = recent.get("title")
            st.session_state.chat_history = recent.get("messages", [])
        else:
            st.session_state.session_id = None
            st.session_state.session_title = None
            st.session_state.chat_history = []
        st.session_state.session_loaded = True

    # Sidebar
    with st.sidebar:
        st.header("📂 Document Management")
        # Upload category selection
        st.session_state.upload_category = st.radio(
            "Upload to knowledge base:",
            options=["TEXTBOOKs", "LECTURE SLIDEs"],
            index=0 if st.session_state.get("upload_category") == "TEXTBOOKs" else 1,
            horizontal=True,
        )
        uploaded_files = st.file_uploader(
            "Upload Documents",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="Each file must be under 25 MB. Supported formats: PDF, DOCX, TXT."
        )
        max_size = 25 * 1024 * 1024 #25MB max file size for each file
        if uploaded_files:
            small_enough = []
            for f in uploaded_files:
                if f.size > max_size:
                    st.error(f"{f.name} exceeds 25MB size limit and will be skipped.")
                else:
                    small_enough.append(f)
            uploaded_files = small_enough
            # Deduplicate by filename
            seen = set(); unique_files = []
            for f in uploaded_files:
                if f.name not in seen:
                    unique_files.append(f); seen.add(f.name)
            uploaded_files = unique_files
        if uploaded_files and st.button("🚀 Process Documents", use_container_width=True):
            process_documents(uploaded_files, st.session_state.upload_category)

        st.divider()
        _render_kb_listing()

        with st.expander("🔧 Advanced Settings"):
            # keep these in session via return values below (read inside chat section)
            st.session_state.setdefault("show_sources", True)
            st.session_state.setdefault("show_context", False)
            st.session_state.show_sources = st.checkbox("Show sources", value=st.session_state.show_sources)
            st.session_state.show_context = st.checkbox("Show context", value=st.session_state.show_context)
            st.session_state.query_kb = st.radio(
                "Use knowledge base:",
                options=["TEXTBOOKs", "LECTURE SLIDEs", "BOTH"],
                index=(
                    0 if st.session_state.get("query_kb") == "TEXTBOOKs" else 
                    (1 if st.session_state.get("query_kb") == "LECTURE SLIDEs" else 2)
                ),
                horizontal=True,
                help="Choose which vector DB to retrieve from when answering."
            )

        if st.button("🗑 Clear Knowledge Base", use_container_width=True):
            clear_knowledge_base()


    # ---------- Normal Chat Layout (ChatGPT-like) ----------
    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader("💬 Chat")

        # Render existing conversation sequentially
        for entry in st.session_state.chat_history:
            # Support legacy Q/A entries as well as new role-based messages
            if isinstance(entry, dict) and "role" in entry:
                role = entry.get("role", "assistant")
                content = entry.get("content", "")
                with st.chat_message(role):
                    st.markdown(content)
                    # Render optional sources/context if present on assistant messages
                    if role == "assistant":
                        if st.session_state.get("show_sources", True) and entry.get("sources"):
                            st.markdown("**Sources**")
                            for i, source in enumerate(entry.get("sources", []), 1):
                                page = source.get('page_number')
                                url = source.get('url')
                                title = f"📄 Source {i}: {source.get('filename', 'unknown')}{f' — page {page}' if page else ''}"
                                with st.expander(title):
                                    st.text(source.get('content_preview', ''))
                                    if url:
                                        st.markdown(f"[Open in viewer (page {page})]({url})", unsafe_allow_html=True)
                        if st.session_state.get("show_context", False) and entry.get("context"):
                            with st.expander("🔍 Retrieved Context"):
                                st.text_area("Context used for answering:", entry.get("context", ""), height=200)
            elif isinstance(entry, dict) and "question" in entry and "answer" in entry:
                # Legacy record: render as two chat bubbles
                with st.chat_message("user"):
                    st.markdown(entry.get("question", ""))
                with st.chat_message("assistant"):
                    st.markdown(entry.get("answer", ""))
                    if st.session_state.get("show_sources", True) and entry.get("sources"):
                        st.markdown("**Sources**")
                        for i, source in enumerate(entry.get("sources", []), 1):
                            page = source.get('page_number')
                            title = f"📄 Source {i}: {source.get('filename', 'unknown')}{f' — page {page}' if page else ''}"
                            with st.expander(title):
                                st.text(source.get('content_preview', ''))
            else:
                # Unknown format fallback
                with st.chat_message("assistant"):
                    st.markdown(str(entry))

        # Chat input at bottom
        prompt = st.chat_input("Ask Anything related to Network Security Course...")
        if prompt is not None:
            # Instantly show user message in chat window
            with st.chat_message("user"):
                st.markdown(prompt)
            st.session_state.chat_history.append({"role": "user", "content": prompt})

            # Add temporary assistant message for 'Assistant is thinking...'
            st.session_state.chat_history.append({
                "role": "assistant",
                "content": "Assistant is thinking...",
                "sources": [],
                "context": "",
                "kb_used": "TEMP"
            })
            
            


            # Persist after user message and temp assistant message
            if st.session_state.get("session_id"):
                try:
                    save_session(st.session_state.session_id, st.session_state.chat_history, st.session_state.get("session_title"))
                except Exception:
                    pass
            else:
                try:
                    import uuid as _uuid
                    from datetime import datetime as _dt
                    now = _dt.now()
                    time_str = now.strftime("%H:%M")
                    date_str = now.strftime("%Y-%m-%d")
                    new_sid = str(_uuid.uuid4())
                    short_sid = new_sid[:8]
                    title = f"{time_str} {date_str} — {short_sid}"
                    save_session(new_sid, st.session_state.chat_history, title)
                    st.session_state.session_id = new_sid
                    st.session_state.session_title = title
                except Exception:
                    pass

            # Now process the query (summary or normal)
            summary_result = handle_summarization_request(prompt)
            is_summary = isinstance(summary_result, dict) and ("summary" in summary_result or "error" in summary_result)
            # Summarization: treat as normal assistant message with sources
            if is_summary:
                if "error" in summary_result:
                    with st.chat_message("assistant"):
                        st.error(summary_result["error"])
                else:
                    # Streaming summary from Ollama, match RAG streaming UI
                    with st.chat_message("assistant"):
                        placeholder = st.empty()
                        buffer = ""
                        citation = summary_result["citation"]
                        stream_gen = ask_llm_summary_stream(
                            citation.get("full_text", citation["content_preview"]),
                            citation["page_number"],
                            citation["filename"],
                            is_slide="slide" in prompt.lower()
                        )
                        result_final = None
                        first_token = True
                        for event in stream_gen:
                            if event["type"] == "token":
                                buffer += event["text"]
                                placeholder.markdown(buffer)
                                if first_token:
                                    # Update the temporary assistant message in chat_history
                                    st.session_state.chat_history[-1] = {
                                        "role": "assistant",
                                        "content": buffer,
                                        "sources": [citation],
                                        "context": "",
                                        "kb_used": "SUMMARY"
                                    }
                                    first_token = False
                            elif event["type"] == "final":
                                buffer = event["text"]
                                placeholder.markdown(buffer)
                                result_final = buffer
                            elif event["type"] == "error":
                                st.error(event["text"])
                        # Optionally render sources/citations inline
                        if st.session_state.get("show_sources", True):
                            st.markdown("**Sources**")
                            page = citation.get('page_number')
                            url = citation.get('url')
                            title = f"📄 Source: {citation.get('filename', 'unknown')}{f' — page {page}' if page else ''}"
                            with st.expander(title):
                                st.text(citation.get('content_preview', ''))
                                if url:
                                    st.markdown(f"[Open in viewer (page {page})]({url})", unsafe_allow_html=True)
                    # Store as assistant message with sources after streaming is complete
                    st.session_state.chat_history[-1] = {
                        "role": "assistant",
                        "content": result_final or buffer,
                        "sources": [citation],
                        "context": "",
                        "kb_used": "SUMMARY"
                    }
                    if st.session_state.get("session_id"):
                        try:
                            save_session(st.session_state.session_id, st.session_state.chat_history, st.session_state.get("session_title"))
                        except Exception:
                            pass
            else:
                if not st.session_state.documents_processed:
                    st.warning("Please upload and process documents before asking questions.", icon="⚠️")
                elif not prompt.strip():
                    st.warning("Please enter a message.", icon="⚠️")
                elif not (st.session_state.rag_chain_textbooks or st.session_state.rag_chain_slides):
                    st.error("Unexpected error: RAG chains not initialized.")
                else:
                    # Generate assistant reply using RAG context history
                    with st.chat_message("assistant"):
                        with st.spinner("🤔 Thinking..."):
                            selected_kb = st.session_state.get("query_kb")
                            if selected_kb == "BOTH":
                                active_rag = st.session_state.rag_chain_both or st.session_state.rag_chain_textbooks or st.session_state.rag_chain_slides
                            elif selected_kb == "LECTURE SLIDEs":
                                active_rag = st.session_state.rag_chain_slides
                            else:
                                active_rag = st.session_state.rag_chain_textbooks
                            # Stream the response token-by-token (fallback to non-streaming if unavailable)
                            if hasattr(active_rag, "chat_with_context_stream"):
                                placeholder = st.empty()
                                buffer = ""
                                result_final = None
                                first_token = True
                                try:
                                    for event in active_rag.chat_with_context_stream(prompt, chat_history=st.session_state.chat_history):
                                        if event.get("type") == "token":
                                            buffer += event.get("text", "")
                                            placeholder.markdown(buffer)
                                            if first_token:
                                                # Update the temporary assistant message in chat_history
                                                st.session_state.chat_history[-1] = {
                                                    "role": "assistant",
                                                    "content": buffer,
                                                    "sources": event.get("sources", []),
                                                    "context": event.get("context", ""),
                                                    "kb_used": st.session_state.get("query_kb")
                                                }
                                                first_token = False
                                        elif event.get("type") == "final":
                                            result_final = event
                                    result = result_final or {"answer": buffer, "sources": [], "context": ""}
                                except Exception:
                                    # In case streaming fails mid-way, fall back
                                    result = active_rag.chat_with_context(prompt, chat_history=st.session_state.chat_history)
                            else:
                                result = active_rag.chat_with_context(prompt, chat_history=st.session_state.chat_history)
                        answer = result.get("answer", "")
                        # Ensure final rendered answer remains
                        st.markdown(answer)

                        # Update the temporary assistant message in chat_history with the final answer
                        st.session_state.chat_history[-1] = {
                            "role": "assistant",
                            "content": answer,
                            "sources": result.get("sources", []),
                            "context": result.get("context", ""),
                            "kb_used": st.session_state.get("query_kb")
                        }

                        # Persist session after each of the assistant reply
                        if st.session_state.get("session_id"):
                            try:
                                save_session(st.session_state.session_id, st.session_state.chat_history, st.session_state.get("session_title"))
                            except Exception:
                                pass

                        # Optionally render sources/context inline
                        if st.session_state.get("show_sources", True) and result.get("sources", []):
                            # Only show sources with valid filename and content_preview
                            filtered_sources = [s for s in result.get("sources", []) if s.get("filename") and s.get("content_preview")]
                            if filtered_sources:
                                st.markdown("**Sources**")
                                for i, source in enumerate(filtered_sources, 1):
                                    page = source.get('page_number')
                                    url = source.get('url')
                                    title = f"📄 Source {i}: {source.get('filename', 'unknown')}{f' — page {page}' if page else ''}"
                                    with st.expander(title):
                                        st.text(source.get('content_preview', ''))
                                        if url:
                                            st.markdown(f"[Open in viewer (page {page})]({url})", unsafe_allow_html=True)
                        if st.session_state.get("show_context", False) and result.get("context", ""):
                            with st.expander("🔍 Retrieved Context"):
                                st.text_area("Context used for answering:", result.get("context", ""), height=200)

    with col2:
        st.subheader("⏱ Chat History")

        sessions_meta = list_sessions()
        session_options = [
            f"{s.get('title', 'Untitled')} ({s.get('message_count', 0)} msgs) — {s.get('session_id')[:8]}"
            for s in sessions_meta
        ]
        current_sid = st.session_state.get("session_id")
        current_index = 0
        for idx, meta in enumerate(sessions_meta):
            if meta.get("session_id") == current_sid:
                current_index = idx
                break

        selected = st.selectbox(
            "Select Chat Session, for new chat session click 'New Chat' below.",
            options=range(len(sessions_meta)) if sessions_meta else [0],
            format_func=lambda i: session_options[i] if sessions_meta else "(No sessions)",
            index=current_index if sessions_meta else 0,
            disabled=not sessions_meta
        )

        cols = st.columns(2)
        with cols[0]:
            if st.button("🆕 New Chat", use_container_width=True):
                # Do not persist yet; wait until first message to create the session
                st.session_state.session_id = None
                st.session_state.session_title = None
                st.session_state.chat_history = []
                # Prevent immediate sidebar auto-switch back to previously selected session
                st.session_state.ignore_next_session_select = True
                st.rerun()

        with cols[1]:
            if st.button("🗑️ Delete Chat", use_container_width=True, disabled=not sessions_meta):
                sid_to_delete = sessions_meta[selected].get("session_id") if sessions_meta else None
                if sid_to_delete:
                    delete_session(sid_to_delete)
                    if sid_to_delete == current_sid:
                        recent = most_recent_session()
                        if recent:
                            st.session_state.session_id = recent.get("session_id")
                            st.session_state.session_title = recent.get("title")
                            st.session_state.chat_history = recent.get("messages", [])
                        else:
                            # No sessions left; clear current and wait for first message to create a new session
                            st.session_state.session_id = None
                            st.session_state.session_title = None
                            st.session_state.chat_history = []
                    st.rerun()

        if sessions_meta:
            selected_sid = sessions_meta[selected].get("session_id")
            if selected_sid and selected_sid != current_sid and not st.session_state.get("ignore_next_session_select"):
                loaded = load_session(selected_sid)
                if loaded:
                    st.session_state.session_id = loaded.get("session_id")
                    st.session_state.session_title = loaded.get("title")
                    st.session_state.chat_history = loaded.get("messages", [])
                    st.rerun()
        # Reset the ignore flag after rendering the sidebar to allow future switches
        if st.session_state.get("ignore_next_session_select"):
            st.session_state.ignore_next_session_select = False

        # --- Optional Viewer Section (if file selected) ---
        active_file = st.session_state.get("active_view_file")
        if active_file and os.path.exists(active_file):
            file_name = os.path.basename(active_file)
            ext = os.path.splitext(file_name)[1].lower()
            st.subheader(f"👁️ Viewer: {file_name}")
            st.info(f"Copy and open this file path in your browser or PDF reader:")
            st.code(os.path.abspath(active_file), language="text")


# ------------------ Core Functions ------------------
def process_documents(uploaded_files, category: str):
    try:
        with st.spinner("⏳ Processing documents..."):
            # Choose directories and vector store by category
            if category == "LECTURE SLIDEs":
                data_dir = getattr(Config, "DATA_DIR_SLIDES", Config.DATA_DIR)
                vstore = st.session_state.vectorstore_slides
            else:
                data_dir = getattr(Config, "DATA_DIR_TEXTBOOKS", Config.DATA_DIR)
                vstore = st.session_state.vectorstore_textbooks

            os.makedirs(data_dir, exist_ok=True)
            file_paths = []
            for uploaded_file in uploaded_files:
                file_path = os.path.join(data_dir, uploaded_file.name)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                file_paths.append(file_path)
            processor = DocumentProcessor(
                chunk_size=Config.CHUNK_SIZE,
                chunk_overlap=Config.CHUNK_OVERLAP
            )
            documents = processor.process_multiple_documents(file_paths)
            if documents and vstore is not None:
                vstore.add_documents(documents)
                st.session_state.documents_processed = True
                st.success("✅ Successfully processed document(s). You can now ask questions!", icon="🎉")
            else:
                st.error("❌ No documents were successfully processed.", icon="🚫")
    except Exception as e:
        st.error(f"Error processing documents: {str(e)}", icon="🚫")

def ask_question(question: str, show_sources: bool, show_context: bool):
    try:
        # Route to appropriate RAG chain based on selected KB
        selected_kb = st.session_state.get("query_kb")
        if selected_kb == "BOTH":
            rag = st.session_state.rag_chain_both or st.session_state.rag_chain_textbooks or st.session_state.rag_chain_slides
        elif selected_kb == "LECTURE SLIDEs":
            rag = st.session_state.rag_chain_slides
        else:
            rag = st.session_state.rag_chain_textbooks
        if rag is None:
            st.error("Selected knowledge base is not initialized.")
            return
        with st.spinner("🤔 Thinking..."):
            result = rag.chat_with_context(question, chat_history=st.session_state.chat_history)
        st.markdown("### 🧠 Answer")
        st.info(result["answer"])
        st.session_state.chat_history.append({
            "question": question,
            "answer": result["answer"],
            "sources": result.get("sources", [])
        })
        if show_sources and result.get("sources"):
            st.markdown("### 📚 Sources")
            for i, source in enumerate(result["sources"], 1):
                page = source.get('page_number')
                url = source.get('url')
                title = f"📄 Source {i}: {source.get('filename', 'unknown')}{f' — page {page}' if page else ''}"
                with st.expander(title):
                    st.text(source.get('content_preview', ''))
                    if url:
                        st.markdown(f"[Open in viewer (page {page})]({url})", unsafe_allow_html=True)
        if show_context and result.get("context"):
            with st.expander("🔍 Retrieved Context"):
                st.text_area("Context used for answering:", result["context"], height=200)
    except Exception as e:
        st.error(f"Error generating answer: {str(e)}")

def display_chat_history():
    history = st.session_state.get("chat_history", [])
    if not history:
        st.info("No messages yet.")
        return

    # Build user/assistant pairs supporting both legacy and new formats
    pairs = []
    i = 0
    n = len(history)
    while i < n:
        entry = history[i]
        if isinstance(entry, dict) and "question" in entry and "answer" in entry:
            pairs.append((entry.get("question", ""), entry.get("answer", "")))
            i += 1
            continue
        if isinstance(entry, dict) and entry.get("role") == "user":
            user_msg = entry.get("content", "")
            # Pair with next assistant message if present
            if i + 1 < n and isinstance(history[i+1], dict) and history[i+1].get("role") == "assistant":
                assistant_msg = history[i+1].get("content", "")
                pairs.append((user_msg, assistant_msg))
                i += 2
                continue
            else:
                pairs.append((user_msg, ""))
                i += 1
                continue
        i += 1

    # Show last 5 pairs in reverse order (most recent first)
    for q, a in reversed(pairs[-5:]):
        title = (q or a or "(empty)")[:50]
        with st.expander(f"💭 {title}..."):
            if q:
                st.write("**Q:**", q)
            if a:
                st.write("**A:**", a)

def clear_knowledge_base():
    # Delete both vector store collections
    try:
        if st.session_state.vectorstore_textbooks:
            st.session_state.vectorstore_textbooks.delete_collection()
            st.session_state.vectorstore_textbooks = None
            st.session_state.rag_chain_textbooks = None
        if st.session_state.vectorstore_slides:
            st.session_state.vectorstore_slides.delete_collection()
            st.session_state.vectorstore_slides = None
            st.session_state.rag_chain_slides = None
        # Reset combined chain
        st.session_state.rag_chain_both = None
    except Exception as e:
        st.warning(f"Error clearing vector stores: {e}")

    # Remove files in data directories (including legacy DATA_DIR)
    dirs = set([
        getattr(Config, "DATA_DIR_TEXTBOOKS", Config.DATA_DIR),
        getattr(Config, "DATA_DIR_SLIDES", Config.DATA_DIR),
        getattr(Config, "DATA_DIR", "data/documents")
    ])
    for d in dirs:
        try:
            if os.path.isdir(d):
                for name in os.listdir(d):
                    path = os.path.join(d, name)
                    if os.path.isfile(path):
                        try:
                            os.remove(path)
                        except Exception:
                            pass
        except Exception:
            pass

    st.session_state.documents_processed = False
    st.success("🧹 Knowledge base (vectors + files) cleared!")
    st.rerun()

# ------------------ Run App ------------------
if __name__ == "__main__":
    main()
